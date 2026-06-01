import base64
import calendar
import os
import subprocess
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document

from .config import OUTPUT_DIR, TEMPLATE_PATH


MONTHS_FR = {
    1: "janvier",
    2: "fevrier",
    3: "mars",
    4: "avril",
    5: "mai",
    6: "juin",
    7: "juillet",
    8: "aout",
    9: "septembre",
    10: "octobre",
    11: "novembre",
    12: "decembre",
}


@dataclass(frozen=True)
class GeneratedReceipt:
    tenant_name: str
    tenant_email: str
    month_label: str
    docx_path: Path
    pdf_path: Path


def receipt_dates(month: int, year: int) -> dict[str, str]:
    """Return every date label needed by the receipt template."""
    if month < 1 or month > 12:
        raise ValueError("Le mois doit etre compris entre 1 et 12.")

    last_day = calendar.monthrange(year, month)[1]
    return {
        "date_mois_et_annee": f"{MONTHS_FR[month]} {year}",
        "date_premier_jour_du_mois": date(year, month, 1).strftime("%d/%m/%Y"),
        "date_dernier_jour_du_mois": date(year, month, last_day).strftime("%d/%m/%Y"),
        "date_paiement": date(year, month, 5).strftime("%d/%m/%Y"),
    }


def build_template_data(row, month: int, year: int, bailleur_nom: str) -> dict[str, str]:
    """Map one tenant/apartment row to the placeholders used in the Word template."""
    tenant_name = f"{row['prenom'].strip()} {row['nom'].strip()}".strip()
    values = {
        "locataire_prenom": row["prenom"].strip(),
        "locataire_nom": row["nom"].strip(),
        "locataire_nom_prenom": tenant_name,
        "bailleur_nom": bailleur_nom.strip(),
        "civilite": row["civilite"].strip(),
        "locataire_email": row["email"].strip(),
        "appartement_numero_code": row["appartement_numero_code"],
        "appartement_adresse": row["appartement_adresse"],
        "loyer_hors_charge": str(row["loyer_hors_charge"]),
        "loyer_charges_uniquement": str(row["loyer_charges_uniquement"]),
        "loyer_total": str(row["loyer_total"]),
    }
    values.update(receipt_dates(month, year))
    return values


def validate_template_data(data: dict[str, str]) -> None:
    """Fail before document generation if a required placeholder value is empty."""
    missing = [key for key, value in data.items() if value is None or str(value).strip() == ""]
    if missing:
        raise ValueError("Donnees manquantes: " + ", ".join(missing))


def safe_filename(value: str) -> str:
    """Remove characters that cannot safely appear in Windows filenames."""
    forbidden = '<>:"/\\|?*'
    cleaned = "".join("_" if char in forbidden else char for char in value)
    return " ".join(cleaned.split())


def powershell_single_quoted(value: Path) -> str:
    """Escape a filesystem path for a single-quoted PowerShell string."""
    return "'" + str(value).replace("'", "''") + "'"


def fill_docx_template(template_path: Path, output_path: Path, data: dict[str, str]) -> None:
    """Create one DOCX by replacing template placeholders while preserving styling."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    replacements = {"{" + key + "}": str(value) for key, value in data.items()}

    document = Document(template_path)
    replace_in_document(document, replacements)
    document.save(output_path)


def replace_in_document(document: Document, replacements: dict[str, str]) -> None:
    """Replace placeholders in paragraphs and tables."""
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        replace_in_table(table, replacements)


def replace_in_table(table, replacements: dict[str, str]) -> None:
    """Replace placeholders recursively inside Word tables."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                replace_in_table(nested_table, replacements)


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replace placeholders and repair common spacing issues created by Word runs."""
    replace_placeholders_in_runs(paragraph.runs, replacements)
    normalize_generated_spacing(paragraph.runs, replacements)


def replace_placeholders_in_runs(runs, replacements: dict[str, str]) -> None:
    """Replace placeholders even when Word splits them across multiple runs."""
    for run in runs:
        for placeholder, replacement in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

    joined_text = "".join(run.text for run in runs)
    if not any(placeholder in joined_text for placeholder in replacements):
        return

    # Keep the first touched run's style when a placeholder spans several runs.
    while True:
        joined_text = "".join(run.text for run in runs)
        match = next(
            (
                (placeholder, replacement, joined_text.find(placeholder))
                for placeholder, replacement in replacements.items()
                if joined_text.find(placeholder) != -1
            ),
            None,
        )
        if match is None:
            return

        placeholder, replacement, start = match
        end = start + len(placeholder)
        cursor = 0
        first_touched_run = None

        for run in runs:
            run_start = cursor
            run_end = cursor + len(run.text)
            cursor = run_end

            if run_end <= start or run_start >= end:
                continue

            local_start = max(start - run_start, 0)
            local_end = min(end - run_start, len(run.text))
            before = run.text[:local_start]
            after = run.text[local_end:]

            if first_touched_run is None:
                run.text = before + replacement + after
                first_touched_run = run
            else:
                run.text = after


def normalize_generated_spacing(runs, replacements: dict[str, str]) -> None:
    """Repair spacing around replacements when the template has split Word runs."""
    civilite = replacements.get("{civilite}")
    tenant_name = replacements.get("{locataire_nom_prenom}")
    date_start = replacements.get("{date_premier_jour_du_mois}")
    date_end = replacements.get("{date_dernier_jour_du_mois}")

    spacing_replacements = {}
    if civilite and tenant_name:
        spacing_replacements[f"{civilite}{tenant_name}"] = f"{civilite} {tenant_name}"
        spacing_replacements[f"{civilite}  {tenant_name}"] = f"{civilite} {tenant_name}"

    for amount_key in (
        "{loyer_total}",
        "{loyer_hors_charge}",
        "{loyer_charges_uniquement}",
    ):
        amount = replacements.get(amount_key)
        if amount:
            spacing_replacements[f"{amount}euros"] = f"{amount} euros"
            spacing_replacements[f"{amount}  euros"] = f"{amount} euros"

    if date_start:
        spacing_replacements[f"{date_start}au"] = f"{date_start} au"
        spacing_replacements[f"{date_start}  au"] = f"{date_start} au"
    if date_end:
        spacing_replacements[f"{date_end}et"] = f"{date_end} et"
        spacing_replacements[f"{date_end}  et"] = f"{date_end} et"

    if spacing_replacements:
        replace_placeholders_in_runs(runs, spacing_replacements)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert one DOCX to PDF through Microsoft Word on Windows."""
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    ps_script = (
        "$ErrorActionPreference='Stop'; "
        "$word=New-Object -ComObject Word.Application; "
        "$word.Visible=$false; "
        "try { "
        f"$doc=$word.Documents.Open({powershell_single_quoted(docx_path)}); "
        f"$doc.SaveAs([ref]{powershell_single_quoted(pdf_path)}, [ref]17); "
        "$doc.Close($false); "
        "} finally { $word.Quit() }"
    )
    encoded_script = base64.b64encode(ps_script.encode("utf-16le")).decode("ascii")
    command = [
        "powershell",
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-EncodedCommand",
        encoded_script,
    ]
    result = subprocess.run(command, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(
            "Echec de conversion PDF via Microsoft Word.\n"
            + result.stderr.strip()
        )


def generate_receipts(rows, month: int, year: int, bailleur_nom: str) -> list[GeneratedReceipt]:
    """Generate one DOCX and one PDF for each active tenant row."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template introuvable: {TEMPLATE_PATH}")

    generated = []
    month_dir = OUTPUT_DIR / f"{year}-{month:02d}"
    for row in rows:
        data = build_template_data(row, month, year, bailleur_nom)
        validate_template_data(data)

        tenant_name = data["locataire_nom_prenom"]
        filename_base = safe_filename(
            f"Quittance - {tenant_name} - {year}-{month:02d}"
        )
        docx_path = month_dir / f"{filename_base}.docx"
        pdf_path = month_dir / f"{filename_base}.pdf"

        fill_docx_template(TEMPLATE_PATH, docx_path, data)
        convert_docx_to_pdf(docx_path, pdf_path)

        generated.append(
            GeneratedReceipt(
                tenant_name=tenant_name,
                tenant_email=data["locataire_email"],
                month_label=data["date_mois_et_annee"],
                docx_path=docx_path,
                pdf_path=pdf_path,
            )
        )
    return generated


def preview_pdf(pdf_path: Path) -> None:
    """Open a generated PDF for manual review."""
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF introuvable: {pdf_path}")
    os.startfile(pdf_path)
